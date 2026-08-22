"""On-demand issue-document rendering for PMP and procurement artefacts."""

from __future__ import annotations

import json
import re
from io import BytesIO
from typing import Literal

from app.projects.artefact_blocks import strip_block_markers
from app.sitewise.artifact_presentation import issue_export_markdown
from app.sitewise.office_pdf import (
    OfficeConversionError,
    convert_office_to_pdf,
    html_to_pdf_bytes,
)

ExportFormat = Literal["pdf", "docx"]
EXPORT_RENDERER_VERSION = "sitewise-issue-sheet-v4"

# Citation column must fit [n] chips; Status gets the narrative width.
_CONSULTANTS_TABLE_WEIGHTS = (14, 24, 8, 26, 28)

_DECISION_FENCE_RE = re.compile(
    r"```pmp-decision\s*\n(?P<payload>\{.*?\})\s*\n```",
    re.DOTALL,
)


def render_artifact_export(
    markdown: str,
    *,
    export_format: ExportFormat,
    project_title: str,
    artifact_title: str,
    version: int,
    workflow_type: str | None = None,
) -> bytes:
    issue_markdown = _render_static_decisions(
        issue_export_markdown(
            strip_block_markers(markdown),
            project_title=project_title,
        )
    )
    html_body = _markdown_html(issue_markdown)
    document_html = _document_html(
        html_body,
        project_title=project_title,
        artifact_title=artifact_title,
        version=version,
    )
    if export_format == "docx":
        return _docx_bytes(
            html_body,
            project_title=project_title,
            artifact_title=artifact_title,
            version=version,
        )
    if export_format == "pdf":
        try:
            docx = _docx_bytes(
                html_body,
                project_title=project_title,
                artifact_title=artifact_title,
                version=version,
            )
        except Exception:
            return _pdf_bytes(document_html)
        return _pdf_from_office(
            docx,
            filename=f"{_safe_export_stem(artifact_title)}.docx",
            html_fallback=document_html,
        )
    raise ValueError(f"unsupported export format: {export_format}")


def render_workbook_pdf(
    workbook_bytes: bytes,
    *,
    filename: str,
    project_title: str = "Cost Plan",
    version: int = 1,
) -> bytes:
    """Issue the cost-plan workbook as PDF so it matches the Excel download."""
    source_name = filename if filename.lower().endswith(".xlsx") else f"{filename}.xlsx"
    try:
        return convert_office_to_pdf(source_bytes=workbook_bytes, filename=source_name)
    except OfficeConversionError:
        return _pdf_bytes(
            _workbook_preview_html(
                workbook_bytes,
                project_title=project_title,
                version=version,
            )
        )


_THEAD_ROW_RE = re.compile(
    r"<thead>\s*<tr>\s*(?P<cells>(?:<th\b[^>]*>.*?</th>\s*)+)</tr>\s*</thead>",
    re.IGNORECASE | re.DOTALL,
)
_TH_CELL_RE = re.compile(r"<th\b[^>]*>(.*?)</th>", re.IGNORECASE | re.DOTALL)
_IDENTITY_HEADER_LABELS = frozenset(
    {
        "project",
        "project title",
        "client",
        "client / owner",
        "address",
        "site",
        "site / address",
        "site / asset",
        "description",
        "state",
        "taxonomy",
        "budget",
    }
)


def _markdown_html(markdown: str) -> str:
    from markdown_it import MarkdownIt

    html = MarkdownIt("commonmark", {"html": False}).enable("table").render(markdown)
    return _demote_summary_table_headers(html)


def _demote_summary_table_headers(html: str) -> str:
    """Drop Field/... labels; render identity first-rows as body cells, not th."""

    def replace(match: re.Match[str]) -> str:
        cells = _TH_CELL_RE.findall(match.group("cells"))
        if not cells:
            return match.group(0)
        first = re.sub(r"<[^>]+>", "", cells[0]).strip().casefold()
        first = re.sub(r"\s+", " ", first)
        if first == "field":
            return ""
        if first not in _IDENTITY_HEADER_LABELS:
            return match.group(0)
        body_cells = "".join(f"<td>{cell}</td>" for cell in cells)
        return f"<tbody><tr>{body_cells}</tr></tbody>"

    return _THEAD_ROW_RE.sub(replace, html)


def _render_static_decisions(markdown: str) -> str:
    def replace(match: re.Match[str]) -> str:
        try:
            payload = json.loads(match.group("payload"))
        except (json.JSONDecodeError, TypeError):
            return ""
        selected = str(
            payload.get("selected") or payload.get("selected_option") or ""
        ).strip()
        options = payload.get("options")
        selected_label = selected
        if isinstance(options, list):
            for option in options:
                if not isinstance(option, dict):
                    continue
                if str(option.get("value") or option.get("id") or "") == selected:
                    selected_label = str(option.get("label") or selected)
                    break
        label = str(payload.get("label") or payload.get("id") or "Decision")
        rationale = str(payload.get("rationale") or payload.get("source") or "—")
        return (
            "| Decision | Selected position | Basis |\n"
            "| --- | --- | --- |\n"
            f"| {_escape_table(label)} | {_escape_table(selected_label or '—')} | "
            f"{_escape_table(rationale)} |"
        )

    return _DECISION_FENCE_RE.sub(replace, markdown)


def _escape_table(value: str) -> str:
    return " ".join(value.split()).replace("|", "\\|")


def _document_html(
    body: str,
    *,
    project_title: str,
    artifact_title: str,
    version: int,
) -> str:
    return f"""<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>{_html_escape(artifact_title)}</title>
<style>
@page {{ size: A4; margin: 16mm 16mm 18mm; @bottom-right {{ content: "Page " counter(page) " of " counter(pages); color: #5C5F66; font-size: 8pt; }} }}
* {{ box-sizing: border-box; }}
body {{ margin: 0; color: #191C21; background: #fff; font-family: "Hanken Grotesk", Helvetica, Arial, sans-serif; font-size: 9.5pt; line-height: 1.45; }}
.sheet {{ background: #fff; min-height: 100%; padding: 0; }}
.issue-meta {{ align-items: end; border-bottom: 2px solid #000; display: grid; gap: 6mm; grid-template-columns: 1fr auto; margin-bottom: 9mm; padding-bottom: 4mm; }}
.issue-project {{ color: #5C5F66; font-size: 8pt; letter-spacing: .08em; margin: 0 0 1mm; text-transform: uppercase; }}
.issue-version {{ color: #000; font-family: "IBM Plex Mono", monospace; font-size: 8pt; letter-spacing: .08em; text-transform: uppercase; }}
h1 {{ color: #191C21; font-size: 21pt; font-weight: 500; letter-spacing: -.02em; line-height: 1.08; margin: 0 0 8mm; }}
h2 {{ border-bottom: 1px solid #D6D6D0; color: #000; font-size: 13pt; font-weight: 600; break-after: avoid; margin: 8mm 0 3mm; padding-bottom: 1.5mm; }}
h3 {{ color: #191C21; font-size: 10.5pt; font-weight: 600; break-after: avoid; margin: 5mm 0 2mm; }}
p {{ margin: 0 0 3mm; max-width: 52ch; }}
ul, ol {{ margin: 1.5mm 0 3.5mm; padding-left: 6mm; }}
li {{ margin: 0 0 1.2mm; }}
table {{ border-collapse: collapse; font-size: 8.2pt; margin: 3mm 0 5mm; width: 100%; }}
thead {{ display: table-header-group; }}
tr {{ break-inside: avoid; }}
th {{ background: #E8E8E4; color: #000; font-weight: 600; }}
th, td {{ border: 1px solid #D6D6D0; padding: 1.8mm 2mm; text-align: left; vertical-align: middle; }}
td {{ background: #fff; }}
blockquote {{ border-left: 1px solid #000; color: #5C5F66; margin: 3mm 0; padding-left: 4mm; }}
code {{ font-family: "IBM Plex Mono", monospace; font-size: .9em; }}
a {{ color: #000; text-decoration: underline; }}
strong {{ font-weight: 600; }}
</style>
</head>
<body><main class="sheet">
<header class="issue-meta"><p class="issue-project">{_html_escape(project_title)}</p><div class="issue-version">Revision {version}</div></header>
{body}
</main></body></html>"""


def _html_escape(value: str) -> str:
    return (
        value.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _pdf_from_office(
    office_bytes: bytes,
    *,
    filename: str,
    html_fallback: str | None = None,
) -> bytes:
    try:
        return convert_office_to_pdf(source_bytes=office_bytes, filename=filename)
    except (OfficeConversionError, OSError, RuntimeError):
        if html_fallback is None:
            raise
    return _pdf_bytes(html_fallback)


def _pdf_bytes(html: str) -> bytes:
    try:
        from weasyprint import HTML

        return HTML(string=html).write_pdf()
    except Exception:
        return html_to_pdf_bytes(_story_safe_html(html))


def _safe_export_stem(value: str) -> str:
    stem = re.sub(r"[^A-Za-z0-9]+", "_", value).strip("_")
    return stem or "Artefact"


def _story_safe_html(html: str) -> str:
    """Drop WeasyPrint-only @page rules that MuPDF treats as fatal CSS."""
    return re.sub(r"@page\s*\{.*?\}\s*", "", html, count=1, flags=re.DOTALL)


def _workbook_preview_html(
    workbook_bytes: bytes,
    *,
    project_title: str,
    version: int,
) -> str:
    from app.sitewise.cost_plan_workbook import workbook_preview_from_bytes

    preview = workbook_preview_from_bytes(workbook_bytes)
    sections: list[str] = []
    for sheet in preview.sheets:
        rows: list[str] = []
        for row_index, row in enumerate(sheet.rows):
            cell_tag = "th" if row_index == 0 else "td"
            cells = "".join(
                f"<{cell_tag}>{_html_escape(cell)}</{cell_tag}>" for cell in row
            )
            rows.append(f"<tr>{cells}</tr>")
        sections.append(
            f"<h2>{_html_escape(sheet.name)}</h2><table>{''.join(rows)}</table>"
        )
    return _document_html(
        "".join(sections),
        project_title=project_title,
        artifact_title="Cost Plan",
        version=version,
    )


def _consultants_table_weights(rows: list, column_count: int) -> list[int] | None:
    if column_count != len(_CONSULTANTS_TABLE_WEIGHTS) or not rows:
        return None
    header_cells = rows[0].find_all(["th", "td"], recursive=False)
    labels = [cell.get_text(" ", strip=True).casefold() for cell in header_cells]
    if not (
        "discipline" in labels
        and "firm" in labels
        and "fee" in labels
        and "status" in labels
        and "citation" in labels
    ):
        return None
    return list(_CONSULTANTS_TABLE_WEIGHTS)


def _docx_bytes(
    html_body: str,
    *,
    project_title: str,
    artifact_title: str,
    version: int,
) -> bytes:
    from bs4 import BeautifulSoup, NavigableString, Tag
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt, RGBColor

    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(16)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(16)
    section.right_margin = Mm(16)
    section.start_type = WD_SECTION.NEW_PAGE

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor(0x19, 0x1C, 0x21)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Arial")
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Arial")
    for name, size, color, before, after in (
        ("Title", 21, RGBColor(0x19, 0x1C, 0x21), 0, 22),
        ("Heading 1", 13, RGBColor(0x00, 0x00, 0x00), 18, 7),
        ("Heading 2", 10.5, RGBColor(0x19, 0x1C, 0x21), 11, 5),
        ("Heading 3", 9.5, RGBColor(0x19, 0x1C, 0x21), 8, 4),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = name != "Title"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Arial")
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Arial")
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(9.5)
        style.paragraph_format.left_indent = Mm(6)
        style.paragraph_format.first_line_indent = Mm(-3)
        style.paragraph_format.space_after = Pt(2.5)
        style.paragraph_format.line_spacing = 1.12

    document.core_properties.title = artifact_title
    document.core_properties.subject = project_title

    header = section.header.paragraphs[0]
    header.text = f"{project_title}  |  {artifact_title}  |  Revision {version}"
    header.style = styles["Caption"]
    header.paragraph_format.space_after = Pt(4)
    header_run = header.runs[0]
    header_run.font.name = "Arial"
    header_run.font.size = Pt(8)
    header_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    header_border = OxmlElement("w:pBdr")
    header_bottom = OxmlElement("w:bottom")
    header_bottom.set(qn("w:val"), "single")
    header_bottom.set(qn("w:sz"), "8")
    header_bottom.set(qn("w:color"), "000000")
    header_border.append(header_bottom)
    header._p.get_or_add_pPr().append(header_border)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    footer_run = footer.add_run("Page ")
    footer_run.font.name = "Arial"
    footer_run.font.size = Pt(8)
    for instruction, suffix in (("PAGE", " of "), ("NUMPAGES", "")):
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), instruction)
        footer._p.append(field)
        if suffix:
            run = footer.add_run(suffix)
            run.font.name = "Arial"
            run.font.size = Pt(8)

    soup = BeautifulSoup(html_body, "html.parser")

    def add_inline(paragraph, node) -> None:
        if isinstance(node, NavigableString):
            paragraph.add_run(str(node))
            return
        if not isinstance(node, Tag):
            return
        if node.name == "br":
            paragraph.add_run().add_break(WD_BREAK.LINE)
            return
        start = len(paragraph.runs)
        for child in node.children:
            add_inline(paragraph, child)
        for run in paragraph.runs[start:]:
            if node.name in {"strong", "b"}:
                run.bold = True
            elif node.name in {"em", "i"}:
                run.italic = True
            elif node.name == "code":
                run.font.name = "Courier New"
            elif node.name == "a":
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                run.underline = True

    def add_list(tag: Tag, *, ordered: bool) -> None:
        style = "List Number" if ordered else "List Bullet"
        for item in tag.find_all("li", recursive=False):
            paragraph = document.add_paragraph(style=style)
            for child in item.children:
                if isinstance(child, Tag) and child.name in {"ul", "ol"}:
                    continue
                add_inline(paragraph, child)
            for nested in item.find_all(["ul", "ol"], recursive=False):
                add_list(nested, ordered=nested.name == "ol")

    def set_cell_margins(cell) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = tc_pr.first_child_found_in("w:tcMar")
        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)
        for edge, value in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
            element = tc_mar.find(qn(f"w:{edge}"))
            if element is None:
                element = OxmlElement(f"w:{edge}")
                tc_mar.append(element)
            element.set(qn("w:w"), str(value))
            element.set(qn("w:type"), "dxa")

    def set_table_geometry(table, widths: list[int]) -> None:
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        total_width = sum(widths)
        tbl_pr = table._tbl.tblPr
        tbl_width = tbl_pr.first_child_found_in("w:tblW")
        if tbl_width is None:
            tbl_width = OxmlElement("w:tblW")
            tbl_pr.insert(0, tbl_width)
        tbl_width.set(qn("w:w"), str(total_width))
        tbl_width.set(qn("w:type"), "dxa")
        layout = tbl_pr.first_child_found_in("w:tblLayout")
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            tbl_pr.append(layout)
        layout.set(qn("w:type"), "fixed")
        borders = tbl_pr.first_child_found_in("w:tblBorders")
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            tbl_pr.append(borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = borders.find(qn(f"w:{edge}"))
            if border is None:
                border = OxmlElement(f"w:{edge}")
                borders.append(border)
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:color"), "D6D6D0")
        tbl_look = tbl_pr.first_child_found_in("w:tblLook")
        if tbl_look is None:
            tbl_look = OxmlElement("w:tblLook")
            tbl_pr.append(tbl_look)
        tbl_look.set(qn("w:firstRow"), "1")
        tbl_look.set(qn("w:lastRow"), "0")
        tbl_look.set(qn("w:firstColumn"), "0")
        tbl_look.set(qn("w:lastColumn"), "0")
        tbl_look.set(qn("w:noHBand"), "1")
        tbl_look.set(qn("w:noVBand"), "1")
        for grid_column, width in zip(
            table._tbl.tblGrid.gridCol_lst, widths, strict=True
        ):
            grid_column.set(qn("w:w"), str(width))
        for row in table.rows:
            for cell, width in zip(row.cells, widths, strict=True):
                cell_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
                cell_width.set(qn("w:w"), str(width))
                cell_width.set(qn("w:type"), "dxa")
                set_cell_margins(cell)

    def table_widths(rows: list[Tag], column_count: int) -> list[int]:
        usable_twips = int(
            (section.page_width - section.left_margin - section.right_margin) / 635
        )
        consultants_weights = _consultants_table_weights(rows, column_count)
        if consultants_weights is not None:
            weight_total = sum(consultants_weights) or 1
            widths = [
                usable_twips * weight // weight_total for weight in consultants_weights
            ]
            widths[-1] += usable_twips - sum(widths)
            return widths
        minimum = min(720, usable_twips // max(column_count * 2, 1))
        scores = [8] * column_count
        for row in rows:
            for index, cell in enumerate(row.find_all(["th", "td"], recursive=False)):
                scores[index] = max(
                    scores[index], min(len(cell.get_text(" ", strip=True)), 42)
                )
        remaining = max(0, usable_twips - minimum * column_count)
        score_total = sum(scores) or 1
        widths = [minimum + remaining * score // score_total for score in scores]
        widths[-1] += usable_twips - sum(widths)
        return widths

    def add_table(tag: Tag) -> None:
        rows = tag.find_all("tr")
        if not rows:
            return
        column_count = max(
            len(row.find_all(["th", "td"], recursive=False)) for row in rows
        )
        table = document.add_table(rows=0, cols=column_count)
        table.style = "Table Grid"
        for row_index, row_tag in enumerate(rows):
            row = table.add_row()
            cells = row_tag.find_all(["th", "td"], recursive=False)
            for column_index, cell_tag in enumerate(cells):
                cell = row.cells[column_index]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for child in cell_tag.children:
                    add_inline(paragraph, child)
                shading = OxmlElement("w:shd")
                if cell_tag.name == "th" or row_index == 0:
                    shading.set(qn("w:fill"), "E8E8E4")
                    cell._tc.get_or_add_tcPr().append(shading)
                    for run in paragraph.runs:
                        run.bold = True
                else:
                    shading.set(qn("w:fill"), "FFFFFF")
                    cell._tc.get_or_add_tcPr().append(shading)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8.2)
                    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            if row_index == 0:
                header_property = OxmlElement("w:tblHeader")
                header_property.set(qn("w:val"), "true")
                row._tr.get_or_add_trPr().append(header_property)
        set_table_geometry(table, table_widths(rows, column_count))
        document.add_paragraph()

    for node in soup.children:
        if isinstance(node, NavigableString) and not str(node).strip():
            continue
        if not isinstance(node, Tag):
            continue
        if node.name == "h1":
            paragraph = document.add_paragraph(style="Title")
            add_inline(paragraph, node)
        elif node.name == "h2":
            paragraph = document.add_paragraph(style="Heading 1")
            add_inline(paragraph, node)
            paragraph_border = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:color"), "D6D6D0")
            paragraph_border.append(bottom)
            paragraph._p.get_or_add_pPr().append(paragraph_border)
        elif node.name == "h3":
            paragraph = document.add_paragraph(style="Heading 2")
            add_inline(paragraph, node)
        elif node.name == "h4":
            paragraph = document.add_paragraph(style="Heading 3")
            add_inline(paragraph, node)
        elif node.name == "p":
            paragraph = document.add_paragraph()
            add_inline(paragraph, node)
        elif node.name in {"ul", "ol"}:
            add_list(node, ordered=node.name == "ol")
        elif node.name == "table":
            add_table(node)
        elif node.name == "blockquote":
            paragraph = document.add_paragraph(style="Quote")
            add_inline(paragraph, node)
        elif node.name == "pre":
            paragraph = document.add_paragraph()
            run = paragraph.add_run(node.get_text())
            run.font.name = "Courier New"
            run.font.size = Pt(8)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
