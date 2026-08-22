from io import BytesIO

import fitz
import pytest
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import RGBColor
from markdown_it import MarkdownIt

from app.projects.artefact_blocks import (
    materialize_block_identity,
    strip_block_markers,
)
from app.sitewise import artifact_exports
from app.sitewise.artifact_exports import (
    _consultants_table_weights,
    render_artifact_export,
    render_workbook_pdf,
)

MARKDOWN = """# Project Management Plan

## Snapshot

The project brief sets the issued scope. [1]

```pmp-decision
{"id":"delivery","label":"Delivery model","selected":"traditional","options":[{"value":"traditional","label":"Traditional tender"}],"rationale":"Client direction"}
```

## Citation key

[1] `brief.pdf`

## Trace & QA

**Inputs to resolve**
- Tender date
"""


def _document_text(payload: bytes) -> str:
    document = Document(BytesIO(payload))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    headers = [
        paragraph.text
        for section in document.sections
        for paragraph in section.header.paragraphs
    ]
    cells = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    return "\n".join([*headers, *paragraphs, *cells])


def test_docx_export_keeps_issue_content_and_selected_decisions_without_trace() -> None:
    payload = render_artifact_export(
        MARKDOWN,
        export_format="docx",
        project_title="Walsh Renovation",
        artifact_title="Project Management Plan",
        version=3,
    )

    assert payload.startswith(b"PK")
    text = _document_text(payload)
    assert "Walsh Renovation" in text
    assert "Traditional tender" in text
    assert "Citation key" in text
    assert "brief.pdf" in text
    assert "Trace & QA" not in text
    assert "Tender date" not in text

    document = Document(BytesIO(payload))
    table = document.tables[0]
    layout = table._tbl.tblPr.first_child_found_in("w:tblLayout")
    assert layout is not None
    assert layout.get(qn("w:type")) == "fixed"
    assert all(
        int(column.get(qn("w:w"))) > 0 for column in table._tbl.tblGrid.gridCol_lst
    )
    assert table.rows[0]._tr.get_or_add_trPr().find(qn("w:tblHeader")) is not None
    assert "NUMPAGES" in document.sections[0].footer._element.xml


def _cell_fill(cell) -> str | None:
    properties = cell._tc.tcPr
    if properties is None:
        return None
    shading = properties.find(qn("w:shd"))
    if shading is None:
        return None
    return shading.get(qn("w:fill"))


def test_docx_export_uses_plain_white_tables_and_black_type() -> None:
    payload = render_artifact_export(
        """# Project Management Plan

## Snapshot

See [the brief](https://example.com).

| Item | Value |
| --- | --- |
| A | One |
| B | Two |
| C | Three |
""",
        export_format="docx",
        project_title="Walsh Renovation",
        artifact_title="Project Management Plan",
        version=3,
    )

    document = Document(BytesIO(payload))
    table = document.tables[0]
    header_fills = [_cell_fill(cell) for cell in table.rows[0].cells]
    body_fills = [
        _cell_fill(cell) for row in table.rows[1:] for cell in row.cells
    ]
    assert header_fills == ["E8E8E4", "E8E8E4"]
    assert body_fills == ["FFFFFF"] * 6
    assert all(
        cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for row in table.rows
        for cell in row.cells
    )
    assert all(
        (paragraph.paragraph_format.space_after or 0) == 0
        and (paragraph.paragraph_format.space_before or 0) == 0
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )
    header_colors = [
        run.font.color.rgb
        for cell in table.rows[0].cells
        for paragraph in cell.paragraphs
        for run in paragraph.runs
        if run.text.strip()
    ]
    assert header_colors
    assert all(color == RGBColor(0x00, 0x00, 0x00) for color in header_colors)
    assert document.styles["Heading 1"].font.color.rgb == RGBColor(0x00, 0x00, 0x00)
    header_run = document.sections[0].header.paragraphs[0].runs[0]
    assert header_run.font.color.rgb == RGBColor(0x00, 0x00, 0x00)
    link_colors = [
        run.font.color.rgb
        for paragraph in document.paragraphs
        for run in paragraph.runs
        if run.underline
    ]
    assert link_colors
    assert all(color == RGBColor(0x00, 0x00, 0x00) for color in link_colors)
    html = artifact_exports._document_html(
        "<h2>Snapshot</h2><table><thead><tr><th>Item</th></tr></thead>"
        "<tbody><tr><td>A</td></tr></tbody></table>",
        project_title="Walsh Renovation",
        artifact_title="Project Management Plan",
        version=3,
    )
    assert "2F72C4" not in html
    assert "2C3037" not in html
    assert "nth-child(even)" not in html
    assert "vertical-align: middle" in html
    assert "background: #E8E8E4" in html


@pytest.mark.parametrize(
    (
        "artifact_title",
        "workflow_type",
        "section_title",
        "header_cells",
        "row_cells",
    ),
    [
        (
            "Project Management Plan",
            "create_pmp",
            "Project Summary",
            ("Project", "Walsh 2"),
            ("Address", "42 Hargrave Street, Paddington"),
        ),
        (
            "Request for Proposal - Structural engineer",
            "consultant_procurement_structural_engineer",
            "Information register",
            ("Document number", "Title"),
            ("420", "Structural details"),
        ),
        (
            "Request for Tender - Mechanical",
            "trade_rft_mechanical",
            "Price schedule",
            ("Price item", "Amount"),
            ("Tender total", "$125,000"),
        ),
    ],
    ids=("pmp", "rfp", "rft"),
)
def test_materialized_tables_round_trip_through_gfm_and_docx_export(
    artifact_title: str,
    workflow_type: str,
    section_title: str,
    header_cells: tuple[str, str],
    row_cells: tuple[str, str],
) -> None:
    source = (
        f"# {artifact_title}\n\n"
        f"## {section_title}\n\n"
        f"| {header_cells[0]} | {header_cells[1]} |\n"
        "| --- | --- |\n"
        f"| {row_cells[0]} | {row_cells[1]} |\n"
    )
    stamped = materialize_block_identity(source, actor_source="ai")

    assert strip_block_markers(stamped.markdown) == source
    html = MarkdownIt("commonmark", {"html": False}).enable("table").render(
        stamped.markdown
    )
    assert html.count("<table>") == 1
    assert html.count("<th>") == 2
    assert html.count("<td>") == 2
    assert "clerk:block" not in html

    payload = render_artifact_export(
        stamped.markdown,
        export_format="docx",
        project_title="Walsh 2",
        artifact_title=artifact_title,
        version=1,
        workflow_type=workflow_type,
    )

    document = Document(BytesIO(payload))
    assert len(document.tables) == 1
    text = _document_text(payload)
    assert header_cells[1] in text
    assert row_cells[0] in text
    assert row_cells[1] in text
    assert "clerk:block" not in text


def test_pdf_export_uses_the_same_docx_bytes(monkeypatch) -> None:
    captured: dict[str, bytes | str] = {}

    def fake_convert(*, source_bytes: bytes, filename: str) -> bytes:
        captured["filename"] = filename
        captured["source"] = source_bytes
        return b"%PDF-from-docx"

    monkeypatch.setattr(artifact_exports, "convert_office_to_pdf", fake_convert)

    payload = render_artifact_export(
        MARKDOWN,
        export_format="pdf",
        project_title="Walsh Renovation",
        artifact_title="Project Management Plan",
        version=3,
    )

    assert payload == b"%PDF-from-docx"
    assert str(captured["filename"]).endswith(".docx")
    docx = render_artifact_export(
        MARKDOWN,
        export_format="docx",
        project_title="Walsh Renovation",
        artifact_title="Project Management Plan",
        version=3,
    )
    assert captured["source"] == docx


def test_pdf_export_falls_back_when_office_conversion_raises_oserror(
    monkeypatch,
) -> None:
    monkeypatch.setattr(
        artifact_exports,
        "convert_office_to_pdf",
        lambda **_kwargs: (_ for _ in ()).throw(OSError("invalid filename")),
    )

    payload = render_artifact_export(
        MARKDOWN,
        export_format="pdf",
        project_title="Walsh Renovation",
        artifact_title="Request for Tender: Structural engineer",
        version=1,
        workflow_type="consultant_procurement_structural_engineer",
    )

    assert payload.startswith(b"%PDF")
    assert len(payload) > 100


def test_pdf_export_falls_back_when_docx_render_fails(monkeypatch) -> None:
    monkeypatch.setattr(
        artifact_exports,
        "_docx_bytes",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(RuntimeError("docx boom")),
    )

    payload = render_artifact_export(
        MARKDOWN,
        export_format="pdf",
        project_title="Walsh Renovation",
        artifact_title="Request for Tender - Structural engineer",
        version=1,
        workflow_type="consultant_procurement_structural_engineer",
    )

    assert payload.startswith(b"%PDF")
    assert len(payload) > 100


def test_pdf_export_smoke_when_a_native_pdf_renderer_is_available() -> None:
    stamped = materialize_block_identity(MARKDOWN, actor_source="ai")
    try:
        payload = render_artifact_export(
            stamped.markdown,
            export_format="pdf",
            project_title="Walsh Renovation",
            artifact_title="Project Management Plan",
            version=3,
        )
    except (ImportError, OSError, RuntimeError) as exc:
        pytest.skip(str(exc))

    assert payload.startswith(b"%PDF")
    assert len(payload) > 100
    with fitz.open(stream=payload, filetype="pdf") as document:
        text = "\n".join(page.get_text() for page in document)
    assert "The project brief sets the issued scope." in text
    assert "clerk:block" not in text


def test_consultants_table_weights_favour_status_and_citation() -> None:
    class _Cell:
        def __init__(self, text: str) -> None:
            self._text = text

        def get_text(self, *_args, **_kwargs) -> str:
            return self._text

    class _Row:
        def __init__(self, labels: list[str]) -> None:
            self._cells = [_Cell(label) for label in labels]

        def find_all(self, *_args, **_kwargs) -> list[_Cell]:
            return self._cells

    weights = _consultants_table_weights(
        [
            _Row(
                [
                    "Discipline",
                    "Firm",
                    "Fee",
                    "Status",
                    "Citation",
                ]
            )
        ],
        5,
    )

    assert weights == [14, 24, 8, 26, 28]
    assert weights[3] > weights[0]
    assert weights[4] > weights[2]
    assert _consultants_table_weights([_Row(["Item", "Status"])], 2) is None


def test_pdf_renderer_never_receives_internal_block_markers(monkeypatch) -> None:
    source = """# Request for Proposal

## Information register

| Document | Title |
| --- | --- |
| 420 | Structural details |
"""
    stamped = materialize_block_identity(source, actor_source="ai")
    rendered: dict[str, str] = {}

    def capture_pdf(html: str) -> bytes:
        rendered["html"] = html
        return b"%PDF-test"

    monkeypatch.setattr(artifact_exports, "_pdf_bytes", capture_pdf)
    monkeypatch.setattr(
        artifact_exports,
        "convert_office_to_pdf",
        lambda **_kwargs: (_ for _ in ()).throw(
            artifact_exports.OfficeConversionError("soffice missing")
        ),
    )

    payload = render_artifact_export(
        stamped.markdown,
        export_format="pdf",
        project_title="Walsh 2",
        artifact_title="Request for Proposal",
        version=1,
        workflow_type="consultant_procurement_structural_engineer",
    )

    assert payload == b"%PDF-test"
    assert "<table>" in rendered["html"]
    assert "Structural details" in rendered["html"]
    assert "clerk:block" not in rendered["html"]


def test_workbook_pdf_converts_the_excel_bytes(monkeypatch) -> None:
    captured: dict[str, bytes | str] = {}

    def fake_convert(*, source_bytes: bytes, filename: str) -> bytes:
        captured["filename"] = filename
        captured["source"] = source_bytes
        return b"%PDF-xlsx"

    monkeypatch.setattr(artifact_exports, "convert_office_to_pdf", fake_convert)

    payload = render_workbook_pdf(b"xlsx-bytes", filename="Cost_Plan_v02.draft.xlsx")

    assert payload == b"%PDF-xlsx"
    assert captured == {
        "filename": "Cost_Plan_v02.draft.xlsx",
        "source": b"xlsx-bytes",
    }


def test_workbook_pdf_falls_back_to_preview_html(monkeypatch) -> None:
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Summary"
    sheet["A1"] = "Cost Code"
    sheet["B1"] = "Item"
    sheet["A2"] = "FEES"
    sheet["B2"] = "Architect"
    buffer = BytesIO()
    workbook.save(buffer)

    monkeypatch.setattr(
        artifact_exports,
        "convert_office_to_pdf",
        lambda **_kwargs: (_ for _ in ()).throw(
            artifact_exports.OfficeConversionError("soffice missing")
        ),
    )
    rendered: dict[str, str] = {}

    def capture_pdf(html: str) -> bytes:
        rendered["html"] = html
        return b"%PDF-preview"

    monkeypatch.setattr(artifact_exports, "_pdf_bytes", capture_pdf)

    payload = render_workbook_pdf(
        buffer.getvalue(),
        filename="Cost_Plan_v02.draft.xlsx",
        project_title="Walsh 2",
        version=2,
    )

    assert payload == b"%PDF-preview"
    assert "Summary" in rendered["html"]
    assert "Architect" in rendered["html"]
    assert "Walsh 2" in rendered["html"]
